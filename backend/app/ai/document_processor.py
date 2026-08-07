import uuid
import PyPDF2
from docx import Document as DocxDocument
from typing import List
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        import io
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text

    def extract_text_from_docx(self, file_content: bytes) -> str:
        import io
        doc = DocxDocument(io.BytesIO(file_content))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text

    def extract_text_from_txt(self, file_content: bytes) -> str:
        return file_content.decode("utf-8")

    def process_document(self, file_content: bytes, filename: str, metadata: dict) -> List[Document]:
        if filename.endswith(".pdf"):
            text = self.extract_text_from_pdf(file_content)
        elif filename.endswith(".docx"):
            text = self.extract_text_from_docx(file_content)
        elif filename.endswith(".txt"):
            text = self.extract_text_from_txt(file_content)
        else:
            raise ValueError("Unsupported file format")

        chunks = self.text_splitter.split_text(text)
        
        documents = []
        for chunk in chunks:
            # We inject a unique chunk ID and preserve the original metadata
            chunk_metadata = metadata.copy()
            chunk_metadata["chunk_id"] = str(uuid.uuid4())
            
            # Context Enrichment: Inject metadata directly into the chunk string!
            source = metadata.get('source', filename)
            enriched_content = f"[Context from Document: {source}]\n{chunk}"
            
            documents.append(Document(page_content=enriched_content, metadata=chunk_metadata))
            
        return documents

document_processor = DocumentProcessor()
